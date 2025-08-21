import base64
import os
import uuid
import json
import logging
import subprocess
import requests
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import Flask, jsonify, request, render_template, redirect, url_for, flash, session, send_file
from werkzeug.utils import secure_filename
from html import escape
import shutil

app = Flask(__name__, static_folder='static')
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your-secret-key')

# Directory configuration
app.config['DOCX_DIR'] = os.path.join(os.getcwd(), 'temp_docx')
app.config['OUTPUT_DIR'] = os.path.join(os.getcwd(), 'temp_pdf')
app.config['SIGNATURE_DIR'] = os.path.join(os.getcwd(), 'temp_signatures')
app.config['EDIT_HISTORY_DIR'] = os.path.join(os.getcwd(), 'edit_history')

# Ensure directories exist
for directory in [app.config['DOCX_DIR'], app.config['OUTPUT_DIR'], app.config['SIGNATURE_DIR'], app.config['EDIT_HISTORY_DIR']]:
    os.makedirs(directory, exist_ok=True)

# Logging setup
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# LibreOffice path
LIBREOFFICE_PATH = r"program\LibreOffice\program\soffice.exe"

# API Base URL
BASE_URL = os.getenv('API_BASE_URL', 'https://api-ticket-system.chervicaon.com/api/v1')

# Allowed file extensions for signatures
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_date(date_str):
    try:
        datetime.strptime(date_str, '%Y-%m-%d')
        return True
    except ValueError:
        return False

def sanitize_input(input_str):
    if not input_str:
        return ""
    return escape(str(input_str).strip())

def save_signature(data, prefix, is_file=False):
    try:
        filename = f"{prefix}_{uuid.uuid4().hex}.png"
        filepath = os.path.join(app.config['SIGNATURE_DIR'], filename).replace('\\', '/')
        logger.debug(f"Attempting to save signature to: {filepath}")

        if is_file:
            if not allowed_file(data.filename):
                logger.error(f"Invalid file extension for {data.filename}")
                return None
            data.save(filepath)
        else:
            if not data or not isinstance(data, str) or not data.startswith('data:image'):
                logger.error("Invalid canvas signature: Data is empty or does not start with 'data:image'")
                return None
            if ',' not in data:
                logger.error("Invalid canvas signature: No comma found in data URI")
                return None
            try:
                img_data = base64.b64decode(data.split(',')[1])
            except (base64.binascii.Error, IndexError) as e:
                logger.error(f"Invalid base64 data: {e}")
                return None
            with open(filepath, 'wb') as f:
                f.write(img_data)

        if not os.path.exists(filepath):
            logger.error(f"Signature file was not created: {filepath}")
            return None

        logger.info(f"Saved signature to: {filepath}")
        session.setdefault('signatures', {})
        session['signatures'][filename] = filepath
        session.modified = True
        return filename
    except Exception as e:
        logger.error(f"Error saving signature: {e}")
        return None

def add_page_number(section):
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    
    run.text = 'Page '
    
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'PAGE')
    run._r.append(fld_simple)
    
    run = footer_paragraph.add_run()
    run.text = ' of '
    
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'NUMPAGES')
    run._r.append(fld_simple)

def create_document(content):
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            add_page_number(section)

        for sig_key in ['customer_signature', 'chervic_signature']:
            if sig_key in content:
                if not os.path.exists(content[sig_key]):
                    logger.error(f"Signature file not found: {content[sig_key]}")
                    raise FileNotFoundError(f"Signature file not found: {content[sig_key]}")

        company_name = content.get('name', 'Company Name')
        start_date = content.get('start_date', 'Start Date')
        location_headquarters = content.get('headquartersLocation', 'Location / Headquarters')
        business_license_number = content.get('registrationNumber', 'Business License Number')
        billing_address = content.get('billingAddress', 'Billing Address')
        billing_contact_name = content.get('billing_contact_name', 'Billing Contact Name')
        contact_person_designation = content.get('contact_person_designation', 'Contact person Designation')
        contact_person_number = content.get('contact_person_number', 'Contact person number')
        contact_person_signature_date = content.get('contact_person_sign_date', 'Contact person Signature Date')
        cas_signature_date = content.get('chervic_date', 'CAS Signature Date')
        billing_email = content.get('billing_email', 'Billing Email')

        heading = doc.add_heading("Master Services Agreement", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.style.font.name = 'Times New Roman'
        heading.runs[0].font.size = Pt(24)

        subheading = doc.add_heading("Between", level=2)
        subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subheading.runs[0].font.size = Pt(14)

        heading = doc.add_heading(f"{company_name}", level=3)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("And")
        run.font.size = Pt(12)
        run.bold = True

        heading = doc.add_heading("Chervic Advisory Services Private Limited", level=3)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].bold = True

        doc.add_page_break()

        p = doc.add_paragraph(f"THIS MASTER SERVICES AGREEMENT (the “Agreement”) is made and effective from {start_date} by & between:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph(f"{company_name} is a company existing and operating in {location_headquarters}, having Business License Number {business_license_number}, with its place of business {billing_address}")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("And")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph("Chervic Advisory Services Private Limited established under laws governing India, with its registered office at Unit No. 7, Sigma Soft Tech Park, Gamma Block, Ground Floor, Whitefield, Bangalore – 560066, Karnataka, India.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph(f"Chervic Advisory Services Private Limited and {company_name} hereinafter referred to individually as a “Party” and collectively as the “Parties”.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("The Company and Service Provider are hereinafter individually referred to as a 'Party' and collectively as 'Parties'. Capitalized terms used but not defined herein shall have the meaning ascribed to such terms in the Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("Definitions", level=2)

        p = doc.add_paragraph("Under this Agreement:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("1) “Affiliate” shall mean, with respect to any entity, any other entity that owns or controls, is owned or controlled by, or is under common ownership or control with such entity. The Parties acknowledge that Service Provider’s Affiliate may provide Services to Company. In such event, Company and the Service Provider’s Affiliate shall execute a separate SOW for Services. Company’s Affiliates may also obtain Services from Service Provider or Service Provider’s Affiliate under the terms of this Agreement by executing a separate SOW for Services. Such SOW shall be governed by terms and conditions as specified in Part B of this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("2) “Resource” would mean any resource person employed by the Service Provider for the performance of its obligation under this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3) “Party” would mean either the Company or the Service Provider.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("1. Scope", level=2)

        p = doc.add_paragraph("The Company shall engage the Service Provider for the provisions of certain services or deliverables (the “Services”) by issuance of statements of work under the terms of this Agreement (the “SOW”). The SOWs issued under this Agreement shall contain all relevant information such as the commercials, delivery date, scope of services etc.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("2. Intellectual Property", level=2)

        heading = doc.add_heading("2.1 The Service Provider will:", level=3)

        p = doc.add_paragraph("2.1.1 inform the Company of any matter which may come to its/Resource’s notice during the operation of this Agreement which may be of interest or importance or use to the Company; and")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("2.1.2 communicate to the Company any proposals or suggestions occurring to it during the operation of this Agreement which may be of service for the business of the Company.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("2.2 However, Service Provider shall retain all right, title and interest in and to the Service Provider’s pre-existing IP, including all right, title and interest in any modifications, customizations, updates, upgrades, enhancements, alterations, made thereto, whether at the request of Company or otherwise, and feedback related thereto.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("2.3 Any intellectual property (IP) that the Service Provider creates in the course of performing work under this Agreement, including the Statement of Work (SOW) — such as trademarks, copyrights, designs, or any other creative assets — shall remain the sole and exclusive property of the Service Provider. Such IP shall not be considered “work for hire” for the Company. The Company shall not acquire any ownership rights over any IP created by the Service Provider, whether during or after the term of this Agreement, unless otherwise expressly Agreed in writing.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("3. Confidential Information", level=2)

        heading = doc.add_heading("3.1 Definitions", level=3)

        p = doc.add_paragraph("“Confidential Information” means any and all information of any kind whatsoever disclosed by one party (Disclosing Party) or any of its Representatives to the other party (Receiving Party) or any of its Representatives prior to, or after, the date of this Agreement in whatever form including, but not limited to, information or discussions related to any business opportunities and any other information which may reasonably be considered as confidential information in the normal course of business of the disclosing Party including but not limited to processes, strategies, data, know-how, trade secrets, designs, reports, test results, drawings, specifications, technical literature and other information or material whether in oral, written, graphic or electromagnetic form (and including without limitation any notes, information or analyses derived from such information however it is produced);")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("“Representatives” means the directors, officers, employees and consultants of the Receiving Party and its associated companies together with any professional advisors of the Receiving Party which it consults in relation to pursuing business opportunities.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("3.2 Obligations", level=3)

        p = doc.add_paragraph("Confidential Information disclosed by the Disclosing Party to the Receiving Party shall be treated as confidential and safeguarded by the Receiving Party in accordance with this Agreement for a period of 3 years from the date of this Agreement. The Receiving Party agrees with and undertakes to the Disclosing Party that it shall and shall procure that its Representatives shall for a period of 3 years from the date of this Agreement:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.2.1 keep in strict confidence and in safe custody any Confidential Information disclosed to the Receiving Party by the Disclosing Party;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.2.2 not use or exploit any Confidential Information other than in connection with pursuing business opportunities;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.2.3 not copy or reproduce any or all of the Confidential Information except as is reasonably necessary in connection with the discussions on business opportunities;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.2.4 promptly comply with any reasonable directions of the Disclosing Party which are given for the protection of the security of the Confidential Information;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.2.5 except as may be required by any applicable law or regulation or the rules or requirements of any relevant stock exchange or relevant regulatory authority, not distribute, disclose or disseminate Confidential Information to anyone, except its Representatives who have a need to know such Confidential Information for the purpose of pursuing business opportunities; and")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.2.6 inform each such Representative of the restrictions as to confidentiality, use and disclosure of such Confidential Information contained in this Agreement and, to the extent that each such Representative is not already under an appropriate duty of confidentiality, impose upon each such Representative obligations of confidentiality at least equivalent to those set out in this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("3.3 Public Statements", level=3)

        p = doc.add_paragraph("Subject to Article 3.4.5 below, each Party hereby undertakes that it shall not (without the prior consent in writing of the other party) release any press statement or make any other announcement to any third party or make any public statement regarding the existence or content of this Agreement or the discussions contemplated by this Agreement or the identity of the Parties to such discussions.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("3.4 Exceptions", level=3)

        p = doc.add_paragraph("The provisions of this Article shall not apply to Confidential Information which the Receiving Party can show to the Disclosing Party's reasonable satisfaction:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.4.1 was known to the Receiving Party (without obligation to keep the same confidential) at the date of disclosure of the Confidential Information by the Disclosing Party;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.4.2 is after the date of disclosure acquired by the Receiving Party in good faith from an independent third party who is not subject to any obligation of confidentiality in respect of such Confidential Information;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.4.3 in its entirety was at the time of its disclosure in the public knowledge or has become public knowledge during the term of this Agreement otherwise than by reason of the Receiving Party's neglect or breach of the restrictions set out in this Agreement or any agreement between the parties;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.4.4 is independently developed by the Receiving Party without access to any or all of the Confidential Information; or")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.4.5 is required by law, judicial action, the rules or regulations of a recognized stock exchange, government department or agency or other regulatory authority to be disclosed in which event the Receiving Party shall take all reasonable steps to consult and take into account the reasonable requirements of the Disclosing Party in relation to such disclosure.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("3.5 Upon the earlier of (i) the expiration or termination of this Agreement, or (ii) written request by the Disclosing Party, the Receiving Party shall, at the Disclosing Party’s option, promptly return or destroy all Confidential Information, including all copies thereof, in its possession or in the possession of its Representatives, whether in written, graphic, electronic, or any other form capable of return or destruction")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("Such return or destruction shall be completed within fifteen (15) days from the date of expiration, termination, or written request, as applicable. Upon completion, the Receiving Party shall, upon request, provide written certification confirming that all such Confidential Information has been returned or irretrievably destroyed.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("4. Warranties", level=2)

        p = doc.add_paragraph("Each Party represents and warrants to the other Party that:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.1 it has the legal capacity and has taken all necessary corporate action required to empower and authorise it to enter into this Agreement;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.2 this Agreement constitutes valid, binding obligations enforceable against it in accordance with the terms of this Agreement;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.3 all information provided by it to the other Party in relation to the provision and receipt of the Services under this Agreement is true to the best of its knowledge, information and belief;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.4 the execution of this Agreement and the performance of its obligations hereunder does not and shall not:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.4.1 contravene any applicable law;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.4.2 contravene any provision of the Party’s constitutional documents;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.4.3 conflict with, or constitute a breach of any of the provisions of any other agreement, obligation, restriction or undertaking which is binding on the Party; and")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.4.4 no fact or circumstance exists that may impair its ability to comply with all of its obligations in terms of this Agreement;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("4.5 it is not insolvent or unable to pay its debts and has not stopped paying its debts as they fall due.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("The warranties explicitly specified herein are in lieu of all other warranties of any kind, implied, statutory, or in any communication between them, including without limitation, the implied warranties of merchantability, non-infringement, title, and fitness for a particular purpose.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("5. Commencement and Termination of Agreement", level=2)

        heading = doc.add_heading("5.1 Commencement", level=3)

        p = doc.add_paragraph("When executed by both Parties, this Agreement comes into force on the date stated at the head of this Agreement (Effective Date).")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("5.2 Termination", level=3)

        p = doc.add_paragraph("This Agreement shall remain in effect from the date hereof until the earliest to occur of the following:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("5.2.1 Two (2) year from the (Effective Date) of this Agreement;")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("5.2.2 If either Party becomes insolvent or bankrupt, or assigns all or a substantial part of its business or assets for the benefit of its creditor(s), or seized by Receivership or Regulatory Authority, or permits the appointment of a receiver or a receiver and manager for its business or assets, or becomes subject to any judicial, administrative, quasi-Judicial or any other legal proceedings relating to the bankruptcy, insolvency, reorganization or the protection of creditors rights or otherwise ceases to conduct business in the normal course.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("5.2.3 Either Party may terminate this Agreement by providing the other Party with no less than thirty (30) days’ prior written notice of its intention to terminate. Such termination shall be effective only upon mutual written agreement of the Parties")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("5.2.4 If the other Party is in default or commits a material breach of this Agreement (including failure to pay an undisputed amount due hereunder), provide that the aggrieved Party serves a 30-day written notice (a 'Rectification Notice') on the other Party and that Party fails to remedy the breach within that period.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("5.2.5 Termination, completion or cancellation of the last remaining Proposal or Project that the Parties have agreed to pursue under this Agreement; or")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("5.2.6 Either party can terminate this agreement by serving a 90-day notice to other party.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("6. Non-Hire and Non-Solicitation", level=2)

        p = doc.add_paragraph("Neither Party shall actively solicit any of each other’s employee, affiliate, associate, client or independent contractor during the term of the Proposal and for a period of two years following its expiry or earlier termination.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("7. Limitation of Liability", level=2)

        p = doc.add_paragraph("In no event shall the Company be liable for any damages, including but not limited to loss of profits, cost of cover, or other incidental, consequential, or indirect damages, even if the Company has been advised of the possibility of such damages. Similarly, the Service Provider’s liability shall be limited to fees received under this Agreement and shall not include any indirect, incidental, or consequential damages. The Service Provider shall make reasonable efforts to deliver the services in alignment with the timelines, quality standards, and specifications set forth in the applicable Statement of Work (SOW). However, delays or deviations caused by events beyond the Service Provider’s reasonable control (e.g., force majeure events, delays in client dependencies, etc.) shall not constitute a breach of contract and shall not be subject to penalties or termination")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("Penalties for Non-Performance: In the event of a delay or failure in service delivery that is within the Service Provider’s control and is not remedied within ten (10) business days after written notice from the Company, the Service Provider shall be liable to pay a penalty of 0.5% of the total project value per week of delay, subject to a maximum cumulative penalty of 5% of the total project value")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("8. Indemnity", level=2)

        p = doc.add_paragraph("The Service Provider should indemnify and hold harmless the Company against any claims, damages, losses, or expenses arising from their negligence, misconduct, or breach of the Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("Likewise, the Company shall indemnify and hold harmless the Service Provider (Chervic Advisory Services), its officers, employees, and affiliates from and against any claims, damages, losses, or expenses (including reasonable legal fees) arising out of or resulting from the Company’s (WMC’s) negligence, willful misconduct, or breach of this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("9. Security", level=2)

        p = doc.add_paragraph("The Service Provider shall make reasonable efforts to comply with the security-related policies and procedures of the Company’s clients, provided that such policies and procedures are communicated to the Service Provider in writing and in advance. In cases where the client does not have a defined security policy, the Service Provider agrees to follow the Company’s relevant security protocols, to the extent such protocols are reasonable, applicable, and have been clearly communicated in writing prior to the commencement of services.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("The Service Provider shall not be held responsible for non-compliance with any security policy that was not disclosed in writing or that imposes unreasonable or commercially impractical requirements. Any additional compliance obligations outside the scope of this Agreement shall be subject to mutual agreement and may require an amendment to the terms, including potential adjustments in timelines, scope, or fees")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("10. Force Majeure", level=2)

        heading = doc.add_heading("10.1 General", level=3)

        p = doc.add_paragraph("Neither Party will be liable for any delay in performing or for failing to perform their respective obligations to the extent that any such specific delay or failure is caused, directly or indirectly, by an event beyond the reasonable control of the either Party, as the case may be, including fire, flood, earthquake, pandemic, elements of nature, acts of war, terrorism, riots, civil disorders, rebellions or revolutions, change in government policies, strikes, lockouts or labour difficulties, such default or delay, collectively, a “Force Majeure Event”.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("10.2 Notice and Suspension", level=3)

        p = doc.add_paragraph("If, as a result of a Force Majeure Event, it becomes impossible or impractical for any Party to carry out its obligations hereunder in whole or in part, then such obligations shall be suspended to the extent necessary by such Force Majeure Event during its continuance and during such time such Party will not be considered in default or contractual breach provided that the affected Party delivers to the non-affected Party as force majeure Notice.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("10.2.1 The Party affected by such Force Majeure Event (the “Affected Party”) shall give prompt written notice to the other Party (the “Non-Affected Party”) of the nature and probable duration of such Force Majeure Event, the extent of its effects on Affected Party’s performance hereunder, and the steps being taken by the Affected Party to address and remove the Force Majeure Event as soon as reasonably practicable following the onset of the Force Majeure Event (the “Force Majeure Notice”). If the Force Majeure Notice is not delivered within 5 (five) Business Days of the initial occurrence of the Force Majeure Event, then the Force Majeure Event will not be deemed to have occurred until the date on which the Non-Affected Party receives the Force Majeure Notice.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("10.2.2 The provision of “Force Majeure” aforesaid shall not be construed as relieving or waiver to either Party from its obligation under this contract to the other Party to the extent of the performed as well as reasonable performable part.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("10.2.3 In the event that a Force Majeure Event persists for a period exceeding 30 days, the Non – Affected Party may terminate this Agreement and any SOW issued hereunder forthwith with prior notice to the Affected Party.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("10.2.4 Notwithstanding anything stated in this Agreement, a Force Majeure Event shall not affect the liability of the Company to make payments to the Service Provider for Services that have already been rendered by the Service Provider.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("11. Independent Service Provider", level=2)

        p = doc.add_paragraph("Service Provider will remain as an independent Service Provider in its relationship with Company. Nothing in this Agreement shall be deemed to have created a partnership, or joint venture or a contract of employment between Company and Service Provider.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("12. Assignment", level=2)

        p = doc.add_paragraph("The Parties shall not assign, sub-license, mortgage, lien, charge, encumber or otherwise dispose of or transfer this Agreement or any of its rights or obligations under this Agreement without the prior written consent of the other. If either party assigns this Agreement to any third parties, such party shall remain the primary obligor and shall be jointly or severally liable for the performance of its obligations under this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("13. No Waiver", level=2)

        p = doc.add_paragraph("Failure or omission by either Party at any time to enforce or require strict or timely compliance with any provision of this Agreement will not affect or impair that provision, or the right of either Party to avail itself of the remedies it may have in respect of any breach of a provision, in any way. However, nothing agreed aforesaid will prevail over the governing laws.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("14. Severability", level=2)

        p = doc.add_paragraph("Any provision of this Agreement that is or becomes illegal, void or unenforceable will be ineffective to the extent only of such illegality, voidness or unenforceability and will not invalidate the remaining provisions.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("15. Variation", level=2)

        p = doc.add_paragraph("This Agreement may not be changed or modified in any way after it has been signed except in writing signed by or on behalf of all the Parties.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("16. Governing Law and Dispute Resolution", level=2)

        p = doc.add_paragraph("16.1 This Agreement shall be governed by, subject to and construed in accordance with the laws of India.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("16.2 Both parties recognise that occasion may arise when one of the parties may have cause for concern relating to the way in which the other party is meeting its obligations under the terms of this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("16.3 The parties shall each be under a general obligation to use all reasonable endeavours to negotiate in good faith and to settle amicably any dispute of whatever nature arising in connection with this Agreement.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("16.4 If a party considers that a dispute exists it shall notify the other party of the dispute in writing.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("16.5 If after (30) calendar days from the date of raising a dispute notice, any party considers that, despite the good faith efforts of the parties, the dispute is not capable of being settled, the aggrieved party may refer the dispute to the competent court in India. The Indian courts, to the exclusion of all other courts, shall have the jurisdiction to finally settle such dispute.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("17. Authority", level=2)

        p = doc.add_paragraph("Each party hereto represents and warrants that the person executing this Agreement on its behalf has express authority to do so, and in so doing, binds the parties hereto.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("18. Enforcement", level=2)

        p = doc.add_paragraph("This Agreement is enforceable by the original parties to it and by their successors in title and permitted assignees.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("19. Amendment and Extension", level=2)

        p = doc.add_paragraph("This Agreement may be amended, and the Term of this Agreement may be extended prior to its expiry only by an instrument in writing signed by duly authorised representative/s of each of the Parties.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("20. Survival", level=2)

        p = doc.add_paragraph("The termination or expiry of this Agreement shall not affect the obligations of each Party with respect to the provisions as set forth in Articles 3 and 6.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("21. Notices", level=2)

        p = doc.add_paragraph("All notices hereunder shall be given in writing by hand delivery, courier service, or email at the addresses set forth below:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading(f"If to {company_name}", level=3)

        p = doc.add_paragraph(f"{billing_contact_name} \n {contact_person_designation}  \n{company_name}  \n {billing_address} \n {contact_person_number}  \nE-mail:{billing_email}")

        heading = doc.add_heading("If to CHERVIC ADVISORY SERVICES PRIVATE LIMITED", level=3)

        p = doc.add_paragraph("Mr. Vasudevan  \nDirector  \nChervic Advisory Services Private Limited  \nUnit No.7 Sigma Soft Tech Park,  \nGamma Block, Ground Floor,  \nWhitefield, Bangalore – 560066,  \nKarnataka E-mail: accounts@chervic.in")

        heading = doc.add_heading("22. Entire Agreement and Modification", level=2)

        p = doc.add_paragraph(f"22.1 This Agreement contains all terms, conditions and provisions hereof and the entire understandings and all representations of understandings and discussions of the Parties relating thereto. This Agreement supersedes and replaces any and all prior agreements and understandings between {company_name} and CHERVIC ADVISORY SERVICES PRIVATE LIMITED.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph("22.2 All terms and conditions included in this Agreement and its Schedule shall apply to any Project covered under this Agreement, unless mutually modified pursuant to the terms of a Project-Related Appendix.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        heading = doc.add_heading("IN WITNESS WHEREOF", level=2)

        p = doc.add_paragraph("The Parties have caused this Agreement to be signed by their duly authorised representatives and effective the date written first above.")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if content.get('customer_signature') or content.get('chervic_signature'):
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            table.style = 'Table Grid'

            if content.get('customer_signature'):
                cell = table.cell(0, 0)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(f"{company_name}\n{billing_contact_name}\nDesignation: {contact_person_designation}\nDate {contact_person_signature_date}")
                run.add_picture(content['customer_signature'], width=Inches(2))

            if content.get('chervic_signature'):
                cell = table.cell(0, 1)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(f"CHERVIC ADVISORY SERVICES PRIVATE LIMITED\nMr. Vasudevan\nDesignation: Director \nDate: {cas_signature_date}")
                run.add_picture(content['chervic_signature'], width=Inches(2))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info("Successfully created document buffer")
        return buffer
    except Exception as e:
        logger.error(f"Error in create_document: {str(e)}")
        raise

def generate_pdf(content):
    doc_buffer = create_document(content)
    filename = f"MSA_{content['name'].replace(' ', '_')}_{uuid.uuid4().hex}.docx"
    docx_filepath = os.path.join(app.config['DOCX_DIR'], filename).replace('\\', '/')
    pdf_filename = filename.replace('.docx', '.pdf')
    pdf_filepath = os.path.join(app.config['OUTPUT_DIR'], pdf_filename).replace('\\', '/')

    try:
        with open(docx_filepath, 'wb') as f:
            f.write(doc_buffer.getvalue())
        logger.info(f"Saved DOCX file: {docx_filepath}")

        try:
            result = subprocess.run(
                [
                    LIBREOFFICE_PATH,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", app.config['OUTPUT_DIR'],
                    docx_filepath
                ],
                capture_output=True,
                text=True,
                check=True
            )
            logger.info(f"LibreOffice conversion output: {result.stdout}")
            if not os.path.exists(pdf_filepath):
                raise Exception(f"PDF file was not generated at {pdf_filepath}")
        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: stdout={e.stdout}, stderr={e.stderr}")
            raise Exception(f"PDF conversion failed: stdout={e.stdout}, stderr={e.stderr}")
        except Exception as e:
            logger.error(f"Unexpected error during PDF conversion: {str(e)}")
            raise

        session.setdefault('pdfs', {})
        session.setdefault('docxs', {})
        session['pdfs'][pdf_filename] = pdf_filepath
        session['docxs'][pdf_filename] = docx_filepath
        session.modified = True

        with open(pdf_filepath, 'rb') as f:
            pdf_buffer = BytesIO(f.read())

        return pdf_filename, pdf_buffer
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        if os.path.exists(docx_filepath):
            os.unlink(docx_filepath)
        if os.path.exists(pdf_filepath):
            os.unlink(pdf_filepath)
        raise

def save_edit_history(pdf_filename, username, changes):
    history_file = os.path.join(app.config['EDIT_HISTORY_DIR'], f"{pdf_filename}_history.json").replace('\\', '/')
    history_entry = {
        'timestamp': datetime.now().isoformat(),
        'username': username,
        'changes': changes
    }
    try:
        history = []
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                history = json.load(f)
        history.append(history_entry)
        with open(history_file, 'w') as f:
            json.dump(history, f, indent=2)
        logger.info(f"Saved edit history to: {history_file}")
    except Exception as e:
        logger.error(f"Error saving edit history: {e}")

class AuthenticationError(Exception):
    pass

def login_user(username, password):
    logger.debug(f"Starting login for username: {username}")
    try:
        url = f"{BASE_URL}/auth/login"
        headers = {'Content-Type': 'application/json'}
        data = {'username': username, 'password': password}
        response = requests.post(url, json=data, headers=headers)
        response.raise_for_status()
        result = response.json()
        user_info = fetch_user_info(response.cookies)
        user_role = user_info.get('userRole')
        if user_role != 'BUSINESS_DEVELOPMENT_USER':
            logger.warning(f"Login rejected for {username}: Invalid role {user_role}")
            raise AuthenticationError(f"Access denied: User role {user_role} is not authorized.")
        user_data = result.get('user', {})
        user_id = user_data.get('id') or user_data.get('userId') or user_info.get('id') or user_info.get('userId')
        if not user_id:
            logger.error(f"No user ID found in login response or user info for {username}")
            raise AuthenticationError("Authentication failed: User ID not provided by API.")
        session['user'] = {
            'id': user_id,
            'username': user_data.get('username', username),
            'role': user_role
        }
        session['cookies'] = {c.name: c.value for c in response.cookies}
        session.modified = True
        logger.info(f"Login successful for {username} with user_id: {user_id}, role: {user_role}")
        return result
    except requests.RequestException as e:
        logger.error(f"Error during login for {username}: {e}")
        raise AuthenticationError(f"Authentication failed: {str(e)}")

def fetch_user_info(login_cookies=None):
    logger.debug("Fetching user info")
    try:
        url = f"{BASE_URL}/auth/role"
        headers = {'Content-Type': 'application/json'}
        with requests.session() as s:
            if login_cookies:
                s.cookies.update(login_cookies)
            elif 'cookies' in session:
                s.cookies.update(session['cookies'])
            else:
                logger.error("No cookies available for fetch_user_info")
                raise AuthenticationError("No authentication cookies available")
            response = s.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            user_info = response.json()
            logger.debug(f"User info response: {json.dumps(user_info, indent=2)}")
            return user_info
    except requests.RequestException as e:
        logger.error(f"Error fetching user info: {e}")
        raise AuthenticationError(f"Failed to fetch user info: {str(e)}")

def fetch_domain_data_by_ariba(ariba_network_id):
    logger.debug(f"Fetching domain data for Ariba Network ID: {ariba_network_id}")
    try:
        url = f"{BASE_URL}/domain?index=0&limit=10&aribaNetworkId={ariba_network_id}"
        headers = {'Content-Type': 'application/json'}
        with requests.session() as s:
            if 'cookies' in session:
                s.cookies.update(session['cookies'])
            else:
                logger.error("No cookies available for fetch_domain_data_by_ariba")
                raise AuthenticationError("No authentication cookies available")
            response = s.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            domain_list = response.json()
            if not domain_list or not isinstance(domain_list, list) or len(domain_list) == 0:
                logger.warning(f"No domains found for Ariba Network ID: {ariba_network_id}")
                return None
            domain_id = domain_list[0].get('id')
            if not domain_id:
                logger.error(f"No domain ID found in response for Ariba Network ID: {ariba_network_id}")
                return None
            url = f"{BASE_URL}/domain/{domain_id}"
            response = s.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            domain_data = response.json()
            logger.info(f"Successfully fetched domain data for domain ID: {domain_id}")
            return domain_data
    except requests.RequestException as e:
        logger.error(f"Error fetching domain data for Ariba Network ID {ariba_network_id}: {e}")
        flash(f"Error fetching domain data: {str(e)}", "error")
        return None

def fetch_ariba_network_ids(user_id):
    logger.debug(f"Fetching Ariba Network IDs for user_id: {user_id}")
    try:
        url = f"{BASE_URL}/domain?index=0&limit=100"
        headers = {'Content-Type': 'application/json'}
        with requests.session() as s:
            if 'cookies' in session:
                s.cookies.update(session['cookies'])
            else:
                logger.error("No cookies available for fetch_ariba_network_ids")
                raise AuthenticationError("No authentication cookies available")
            response = s.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            domains = response.json()
            ariba_network_ids = [domain.get('aribaNetworkId') for domain in domains if domain.get('aribaNetworkId')]
            logger.info(f"Fetched {len(ariba_network_ids)} Ariba Network IDs")
            return ariba_network_ids
    except requests.RequestException as e:
        logger.error(f"Error fetching Ariba Network IDs: {e}")
        flash(f"Error fetching Ariba Network IDs: {str(e)}", "error")
        return []

@app.route('/fetch_domain_data', methods=['POST'])
def fetch_domain_data():
    if 'user' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    data = request.get_json()
    ariba_network_id = data.get('aribaNetworkId')
    if not ariba_network_id:
        return jsonify({'success': False, 'error': 'Ariba Network ID is required'}), 400
    try:
        domain_data = fetch_domain_data_by_ariba(ariba_network_id)
        if domain_data:
            session['aribaNetworkId'] = ariba_network_id
            session.modified = True
            return jsonify({'success': True, 'domain_data': domain_data})
        else:
            return jsonify({'success': False, 'error': 'No domain data found'}), 404
    except Exception as e:
        logger.error(f"Error fetching domain data: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        try:
            login_user(username, password)
            flash("Login successful!", "success")
            return redirect(url_for('index'))
        except AuthenticationError as e:
            flash(str(e), "error")
            return render_template('login.html')
    return render_template('login.html')

@app.route('/logout')
def logout():
    for filename in session.get('pdfs', {}):
        filepath = session['pdfs'][filename]
        if os.path.exists(filepath):
            try:
                os.unlink(filepath)
                logger.info(f"Deleted PDF file: {filepath}")
            except Exception as e:
                logger.error(f"Error deleting PDF file {filepath}: {e}")

    for filename in session.get('docxs', {}):
        filepath = session['docxs'][filename]
        if os.path.exists(filepath):
            try:
                os.unlink(filepath)
                logger.info(f"Deleted DOCX file: {filepath}")
            except Exception as e:
                logger.error(f"Error deleting DOCX file {filepath}: {e}")

    for filename in session.get('signatures', {}):
        filepath = session['signatures'][filename]
        if os.path.exists(filepath):
            try:
                os.unlink(filepath)
                logger.info(f"Deleted signature file: {filepath}")
            except Exception as e:
                logger.error(f"Error deleting signature file {filepath}: {e}")

    for filename in session.get('edit_history', {}):
        history_file = os.path.join(app.config['EDIT_HISTORY_DIR'], f"{filename}_history.json").replace('\\', '/')
        if os.path.exists(history_file):
            try:
                os.unlink(history_file)
                logger.info(f"Deleted edit history file: {history_file}")
            except Exception as e:
                logger.error(f"Error deleting edit history file {history_file}: {e}")

    directories = [
        app.config['DOCX_DIR'],
        app.config['OUTPUT_DIR'],
        app.config['SIGNATURE_DIR'],
        app.config['EDIT_HISTORY_DIR']
    ]
    for directory in directories:
        if os.path.exists(directory):
            try:
                shutil.rmtree(directory)
                logger.info(f"Deleted directory: {directory}")
            except Exception as e:
                logger.error(f"Error deleting directory {directory}: {e}")

    session.clear()
    flash("Logged out successfully.", "success")
    return redirect(url_for('login'))

@app.route('/')
def index():
    if 'user' not in session:
        return redirect(url_for('login'))
    data = session.get('form_data', {
        "name": "",
        "websiteUrl": "",
        "registrationNumber": "",
        "headquartersLocation": "",
        "countriesOfOperation": "",
        "businessType": "",
        "industryType": "",
        "billingAddress": "",
        "otherLocalTaxId": "",
        "billing_contact_name": "",
        "billing_email": "",
        "start_date": "",
        "contact_person_designation": "",
        "contact_person_number": "",
        "chervic_date": "",
        "contact_person_sign_date": ""
    })
    try:
        user_id = session['user'].get('id')
        if not user_id:
            logger.error("No user_id found in session")
            flash("User ID not found. Please log in again.", "error")
            return redirect(url_for('login'))
        ariba_network_ids = fetch_ariba_network_ids(user_id)
        selected_ariba_id = session.get('aribaNetworkId', '')
        if selected_ariba_id:
            api_data = fetch_domain_data_by_ariba(selected_ariba_id)
            if api_data:
                api_field_mapping = {
                    "name": "name",
                    "websiteUrl": "websiteUrl",
                    "registrationNumber": "registrationNumber",
                    "headquartersLocation": "headquartersLocation",
                    "countriesOfOperation": "countriesOfOperation",
                    "businessType": "businessType",
                    "industryType": "industryType",
                    "billingAddress": "billingAddress",
                }
                for api_key, form_key in api_field_mapping.items():
                    if api_key in api_data and api_data[api_key] is not None:
                        data[form_key] = str(api_data[api_key])
            else:
                logger.warning(f"No domain data found for Ariba Network ID: {selected_ariba_id}")
                flash(f"No domain data found for Ariba Network ID: {selected_ariba_id}. Please enter manually.", "warning")
        session['form_data'] = data
        session.modified = True
        return render_template('index.html', data=data, ariba_network_ids=ariba_network_ids, selected_ariba_id=selected_ariba_id)
    except Exception as e:
        logger.error(f"Error in index route: {e}")
        flash("An error occurred while loading the form. Please try again.", "error")
        return render_template('index.html', data=data, ariba_network_ids=[], selected_ariba_id='')

@app.route('/submit', methods=['POST'])
def generate_nda():
    if 'user' not in session:
        flash("Please log in to continue.", "error")
        return redirect(url_for('login'))
    data = session.get('form_data', {})
    required_fields = [
        "name", "websiteUrl", "registrationNumber", "headquartersLocation",
        "countriesOfOperation", "businessType", "industryType", "billingAddress",
        "billing_contact_name", "billing_email", "start_date",
        "contact_person_designation", "contact_person_number", "chervic_date",
        "contact_person_sign_date"
    ]
    for field in required_fields:
        value = request.form.get(field, '').strip()
        if not value:
            flash(f"{field.replace('_', ' ').title()} is required.", "error")
            return render_template('index.html', data=data)
        data[field] = sanitize_input(value)
    date_fields = ['start_date', 'chervic_date', 'contact_person_sign_date']
    for field in date_fields:
        if not validate_date(data[field]):
            flash(f"Invalid date format for {field.replace('_', ' ').title()}. Use YYYY-MM-DD.", "error")
            return render_template('index.html', data=data)
    optional_fields = ['chervic_name', 'chervic_title', 'customer_sign_name', 'customer_sign_title', 'agreement_date']
    for field in optional_fields:
        data[field] = sanitize_input(request.form.get(field, ''))
    data['aribaNetworkId'] = sanitize_input(request.form.get('aribaNetworkId', session.get('aribaNetworkId', '')))
    if data['aribaNetworkId']:
        session['aribaNetworkId'] = data['aribaNetworkId']
        session.modified = True
    chervic_signature_data = request.form.get('chervic_signature_data')
    chervic_signature_file = request.files.get('chervic_signature')
    if chervic_signature_data and chervic_signature_data.startswith('data:image'):
        signature_filename = save_signature(chervic_signature_data, 'chervic', is_file=False)
        if signature_filename:
            data['chervic_signature'] = session['signatures'][signature_filename]
        else:
            flash("Failed to process Chervic canvas signature.", "error")
            return render_template('index.html', data=data)
    elif chervic_signature_file and allowed_file(chervic_signature_file.filename):
        signature_filename = save_signature(chervic_signature_file, 'chervic', is_file=True)
        if signature_filename:
            data['chervic_signature'] = session['signatures'][signature_filename]
        else:
            flash("Failed to process Chervic file signature.", "error")
            return render_template('index.html', data=data)
    else:
        flash("Chervic signature is required.", "error")
        return render_template('index.html', data=data)
    customer_signature_data = request.form.get('contact_person_signature_data')
    customer_signature_file = request.files.get('contact_person_signature')
    if customer_signature_data and customer_signature_data.startswith('data:image'):
        signature_filename = save_signature(customer_signature_data, 'customer', is_file=False)
        if signature_filename:
            data['customer_signature'] = session['signatures'][signature_filename]
        else:
            flash("Failed to process Customer canvas signature.", "error")
            return render_template('index.html', data=data)
    elif customer_signature_file and allowed_file(customer_signature_file.filename):
        signature_filename = save_signature(customer_signature_file, 'customer', is_file=True)
        if signature_filename:
            data['customer_signature'] = session['signatures'][signature_filename]
        else:
            flash("Failed to process Customer file signature.", "error")
            return render_template('index.html', data=data)
    else:
        flash("Customer signature is required.", "error")
        return render_template('index.html', data=data)
    session['form_data'] = data
    session.modified = True
    try:
        pdf_filename, pdf_buffer = generate_pdf(data)
        username = session['user']['username']
        changes = {
            'fields_updated': {k: v for k, v in data.items() if k in required_fields + optional_fields},
            'signatures_added': {
                'chervic': bool(data.get('chervic_signature')),
                'customer': bool(data.get('customer_signature'))
            }
        }
        save_edit_history(pdf_filename, username, changes)
        flash("Agreement generated successfully!", "success")
        return redirect(url_for('view_pdf', filename=pdf_filename))
    except Exception as e:
        logger.error(f"Error generating Agreement: {e}")
        flash(f"Error generating Agreement: {str(e)}", "error")
        return render_template('index.html', data=data)

@app.route('/download_pdf/<filename>')
def download_pdf(filename):
    if 'user' not in session:
        flash("Please log in to access PDFs.", "error")
        return redirect(url_for('login'))
    filepath = session.get('pdfs', {}).get(filename)
    if not filepath or not os.path.exists(filepath):
        flash("PDF not found.", "error")
        logger.error(f"PDF not found for download: {filepath}")
        return redirect(url_for('index'))
    try:
        return send_file(
            filepath,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"Error downloading PDF file: {e}")
        flash("Error downloading PDF.", "error")
        return redirect(url_for('index'))

@app.route('/download_docx/<filename>')
def download_docx(filename):
    if 'user' not in session:
        flash("Please log in to access documents.", "error")
        return redirect(url_for('login'))
    filepath = session.get('docxs', {}).get(filename)
    if not filepath or not os.path.exists(filepath):
        flash("Document not found.", "error")
        logger.error(f"DOCX not found for download: {filepath}")
        return redirect(url_for('index'))
    try:
        return send_file(
            filepath,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename.replace('.pdf', '.docx')
        )
    except Exception as e:
        logger.error(f"Error downloading DOCX file: {e}")
        flash("Error downloading document.", "error")
        return redirect(url_for('index'))

@app.route('/view_pdf/<filename>')
def view_pdf(filename):
    if 'user' not in session:
        flash("Please log in to view PDFs.", "error")
        return redirect(url_for('login'))
    filepath = session.get('pdfs', {}).get(filename)
    if not filepath or not os.path.exists(filepath):
        logger.error(f"PDF not found: {filepath}")
        flash("PDF not found.", "error")
        return redirect(url_for('index'))
    try:
        pdf_url = url_for('serve_pdf', filename=filename, _external=True)
        logger.info(f"Serving PDF at URL: {pdf_url} for file: {filename}")
        return render_template('view_pdf.html', pdf_filename=filename, pdf_url=pdf_url)
    except Exception as e:
        logger.error(f"Error preparing PDF for viewing: {e}")
        flash("Error loading PDF.", "error")
        return redirect(url_for('index'))

@app.route('/serve_pdf/<filename>')
def serve_pdf(filename):
    if 'user' not in session:
        flash("Please log in to access PDFs.", "error")
        return redirect(url_for('login'))
    filepath = session.get('pdfs', {}).get(filename)
    if not filepath or not os.path.exists(filepath):
        flash("PDF not found.", "error")
        logger.error(f"PDF not found for serving: {filepath}")
        return redirect(url_for('index'))
    try:
        return send_file(
            filepath,
            mimetype='application/pdf',
            as_attachment=False
        )
    except Exception as e:
        logger.error(f"Error serving PDF file: {e}")
        flash("Error serving PDF.", "error")
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)